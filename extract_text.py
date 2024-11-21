import os
import argparse
from docx import Document

# 提取DOCX内容
def extract_text_from_docx(file_path):
    text = ""
    doc = Document(file_path)
    
    # 提取段落文本
    for para in doc.paragraphs:
        if para.text.strip():  # 只添加非空段落
            text += para.text + '\n'
    
    # 提取表格内容
    for table in doc.tables:
        text += "\n=== 表格开始 ===\n"  # 表格标记
        
        # 获取每列的最大宽度
        col_widths = []
        for col in range(len(table.columns)):
            width = max(len(cell.text) for row in table.rows for cell in row.cells[col:col+1])
            col_widths.append(width)
        
        # 输出表格内容
        for row in table.rows:
            # 将每个单元格的文本对齐后用竖线分隔
            row_text = ' | '.join(
                cell.text.ljust(col_widths[i]) 
                for i, cell in enumerate(row.cells)
            )
            text += row_text + '\n'
            
        text += "=== 表格结束 ===\n\n"  # 表格结束标记
    
    return text

# 保存提取的纯文本内容
def save_to_file(text, output_dir, filename):
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

# 提取并保存DOCX中的文本
def extract_and_save(input_path, output_dir):
    docx_files = []
    
    # 收集需要处理的文件
    if os.path.isfile(input_path):
        if not input_path.endswith('.docx'):
            print(f"错误：'{input_path}' 不是DOCX文件")
            return
        docx_files.append(input_path)
    elif os.path.isdir(input_path):
        docx_files.extend([
            os.path.join(input_path, f) 
            for f in os.listdir(input_path) 
            if f.endswith('.docx')
        ])
    else:
        print(f"错误：'{input_path}' 不存在")
        return
    
    total_files = len(docx_files)
    if total_files == 0:
        print("未找到任何DOCX文件")
        return
        
    print(f"\n共发现 {total_files} 个DOCX文件待处理")
    print("=" * 50)
    
    # 处理文件
    for index, file_path in enumerate(docx_files, 1):
        filename = os.path.basename(file_path)
        print(f"\n正在处理第 {index}/{total_files} 个文件: {filename}")
        
        try:
            text = extract_text_from_docx(file_path)
            output_filename = f"{os.path.splitext(filename)[0]}.txt"
            save_to_file(text, output_dir, output_filename)
            print(f"✓ 已完成提取并保存到: {os.path.join(output_dir, output_filename)}")
            
        except Exception as e:
            print(f"✗ 处理文件 '{filename}' 时出错: {str(e)}")
    
    print("\n" + "=" * 50)
    print(f"处理完成！成功提取 {total_files} 个文件的内容到目录: {output_dir}")

def main():
    # 创建命令行参数解析器
    parser = argparse.ArgumentParser(
        description='DOCX文本提取工具 - 可以从Word文档中提取纯文本内容',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
使用示例:
  # 处理单个文件
  python extract_text.py document.docx
  
  # 处理单个文件并指定输出目录
  python extract_text.py document.docx -o /path/to/output
  
  # 处理整个目录下的所有docx文件
  python extract_text.py /path/to/docx/folder
  
  # 处理整个目录并指定输出目录
  python extract_text.py /path/to/docx/folder -o /path/to/output
        '''
    )
    
    parser.add_argument('input', 
                       help='输入路径：可以是单个DOCX文件或包含多个DOCX文件的目录')
    
    parser.add_argument('-o', '--output', 
                       default='output/docx_output',
                       help='输出目录路径，用于存放提取的文本文件 (默认: output/docx_output)')
    
    args = parser.parse_args()
    extract_and_save(args.input, args.output)

if __name__ == "__main__":
    main()
